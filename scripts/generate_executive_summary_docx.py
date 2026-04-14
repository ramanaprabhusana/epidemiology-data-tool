from __future__ import annotations

import argparse
from datetime import date
from pathlib import Path


def _add_bullets(doc, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def build_document(*, doc_date: str):
    from docx import Document
    from docx.shared import Inches

    doc = Document()

    section = doc.sections[0]
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

    doc.add_heading("Executive summary, CLL epidemiology and pipeline (status update)", level=1)
    doc.add_paragraph(f"Date: {doc_date}")

    doc.add_heading("Executive summary", level=2)
    doc.add_paragraph(
        "We have completed the extraction and structuring of core CLL epidemiology, and we have implemented a "
        "hybrid automation data pipeline to consolidate evidence from authoritative sources while maintaining data "
        "quality and full traceability. In parallel, we conducted a focused literature review to identify the "
        "clinical flow parameters required for patient journey modeling, and we established forecasting methods "
        "to project epidemiology over time under conservative and growth scenarios."
    )
    doc.add_paragraph(
        "Our goal is to keep the work easy to check and explain. Every datapoint is tied back to a citation and URL, "
        "and the tables used in modeling keep a clear link between raw evidence, reconciliation decisions, and the "
        "final datasets used for patient flow and forecasting."
    )

    doc.add_heading("What has been completed so far", level=2)
    doc.add_paragraph(
        "Epidemiology extraction and structuring is complete for the core CLL metrics from SEER, including incidence, "
        "prevalence, survival, and mortality, with relevant stratifications. A standardized master dataset has been "
        "created with clear source tagging and data type classification, enabling full traceability from every metric "
        "back to its source. In parallel, the clinical and modeling team completed a literature review to identify "
        "patient journey parameters and defined the calculation logic for derived KPIs, including drug ready population, "
        "treated population, and line of therapy segmentation. Using these inputs, we generated an initial patient flow "
        "from total prevalence through 1L, 2L, and 3L plus pools."
    )
    _add_bullets(
        doc,
        [
            "Core CLL epidemiology extracted and structured from SEER (incidence, prevalence, survival, mortality), with stratifications and source tagging.",
            "Master dataset created with standardized fields and traceability (source tagging and data type classification).",
            "Literature review completed for patient journey parameters, plus defined KPI logic (drug ready, treated, line of therapy segmentation).",
            "Initial patient flow generated from total prevalence through 1L, 2L, and 3L plus pools.",
            "Forecasting implemented using adapted APC (conservative, trend tapered) and log linear (continued growth) approaches.",
            "Population prediction integrated using US Census, with planned extension to UN population inputs for sensitivity assessment.",
        ],
    )

    doc.add_heading("Hybrid automation data pipeline (what we built and why)", level=2)
    doc.add_paragraph(
        "We are designing a hybrid automation data pipeline to address a common problem in oncology epidemiology, "
        "getting access to trusted cancer data from several sources while keeping quality high. Full automation is not "
        "realistic for many portals, so we use three simple modes. First, for sources that allow web access or scraping, "
        "the tool can pull numbers directly and update them more often. Second, for sources that require manual downloads "
        "such as GLOBOCAN, we load structured files and, where helpful, use proxy variables. Third, for desktop based tools, "
        "we plan scheduled manual extractions. We are also testing browser automation for SEER and open APIs, and expect to "
        "run these jobs on free cloud infrastructure (Oracle) where that adds value."
    )
    _add_bullets(
        doc,
        [
            "Source list defined in configuration across gold, silver, and bronze sources (registries, guidelines, journals, and APIs where enabled).",
            "Standardized evidence schema with full context fields (metric, definition, population, year, geography, split logic, value, citation, URL, tier, notes).",
            "Reproducible execution via CLI and notebook, plus standalone CSV only mode for easy handoff and submission.",
            "Exports for stakeholder consumption, including evidence tables, tool ready tables, KPI scorecards, and dashboard ready datasets.",
        ],
    )
    doc.add_paragraph(
        "This setup is meant to be practical for day to day use. It avoids brittle scraping where access is restricted, "
        "and it separates evidence collection from modeling so that updates can be added without redoing the logic. It also "
        "keeps the split between content work and tooling clear, where analysts add evidence in templates and the pipeline "
        "handles validation, scoring, reconciliation, and exports."
    )

    doc.add_heading("Validation framework and quality assurance (procedures and techniques)", level=2)
    doc.add_paragraph(
        "The pipeline uses a two step validation process so that only reliable data moves forward. First, we compare key "
        "metrics across sources and flag cases where values differ more than we expect, so they can be checked. Second, we "
        "apply CLL specific business rules to make sure the numbers are clinically sensible, and we give each record a simple "
        "quality score from 0 to 100 so users can see how much weight to place on it."
    )
    doc.add_paragraph(
        "In practice, we treat the evidence layer as a living repository. When a metric appears inconsistent across sources, the workflow "
        "does not overwrite differences. It surfaces conflicts explicitly, enables reconciliation, and preserves the competing sources so "
        "the client can review the rationale."
    )
    _add_bullets(
        doc,
        [
            "Cross source checks: automatic comparison of key metrics across sources (example, SEER, GLOBOCAN, and internal Roche inputs), with clear flags when they disagree.",
            "CLL specific business rules (examples): mortality cannot exceed incidence, 5 year survival should stay within expected modern era ranges (around 85 to 95 percent), male to female ratios should stay close to 1.5 to 2.0, and older data (more than 4 years old) is marked as stale.",
            "Record level quality scoring: 0 to 100 score based on validation results, completeness, recency, extraction success, and agreement across sources.",
            "Conflict detection and reconciliation outputs, including ranges, recommended values, and readiness status for downstream use.",
        ],
    )

    doc.add_heading("Harmonization and structuring", level=2)
    doc.add_paragraph(
        "Beyond validation, the system smooths out differences in data structure and methods across sources. It standardizes "
        "variable names, makes age adjustment and similar choices explicit, and adds derived metrics on top of the raw evidence. "
        "This lets the modeling team use one consistent structure for patient journey and forecasting work while still keeping "
        "the original sources visible."
    )
    _add_bullets(
        doc,
        [
            "Standardized variable names and consistent output schemas for evidence, tool ready datasets, and scorecards.",
            "Enrichment with derived metrics and scenario selections while retaining raw evidence and citations.",
            "Audit trail outputs, including source logs, validation reports, conflicts, reconciliation, and summary views.",
        ],
    )

    doc.add_heading("Forecasting (approaches and inputs)", level=2)
    doc.add_paragraph(
        "We currently use two simple but transparent forecasting approaches. An adapted APC style approach gives a conservative "
        "view where trends slow over time, and a log linear approach gives a view where growth continues at a similar pace. "
        "Population inputs are based on US Census projections today, and we plan to add UN population projections to test how "
        "sensitive results are to different population baselines and to support more geographies."
    )
    _add_bullets(
        doc,
        [
            "Conservative scenario: adapted APC model with trend tapering assumptions.",
            "Growth scenario: log linear model capturing sustained growth trajectories.",
            "Population inputs: US Census projections used for population prediction, with planned UN based sensitivity runs and comparisons.",
        ],
    )

    doc.add_heading("Ready deliverables", level=2)
    doc.add_paragraph(
        "The following deliverables are ready for review and can be refreshed on demand using the same pipeline run."
    )
    _add_bullets(
        doc,
        [
            "Standardized evidence tables and source logs with tiering and traceability.",
            "KPI scorecard, conflicts table, reconciliation table, and white space coverage summary.",
            "Tool ready outputs and InsightACE epidemiology exports.",
            "Forecast outputs for conservative and growth scenarios.",
            "Runnable notebook and standalone pipeline script for reproducibility, submission, and handoff.",
        ],
    )
    doc.add_paragraph(
        "The automation deliverable also includes a dashboard ready export layer so analytics tools can connect to a stable set of tables. "
        "Where needed, outputs can be consolidated into a single Excel workbook for client review."
    )

    doc.add_heading("Proposed next steps", level=2)
    doc.add_paragraph(
        "In the near term we will focus on checking key assumptions and setting up a simple process for regular updates. In "
        "parallel, we will add modest extensions to the tooling where they clearly reduce manual work, and we will summarize "
        "the approach in a short client facing deck and methods note."
    )
    _add_bullets(
        doc,
        [
            "Validate key modeling assumptions, including treatment uptake, line of therapy transitions, and eligibility definitions.",
            "Agree on how often SEER and literature inputs should be refreshed, and reflect that in a light weight run book.",
            "Add UN population based sensitivity runs and compare outputs versus Census based forecasts.",
            "Finalize a client ready deck and a short methods note (confidence rubric, validation thresholds, reconciliation rules).",
            "Extend the same framework to additional indications using the existing configuration approach.",
        ],
    )

    return doc


def main() -> int:
    parser = argparse.ArgumentParser(description="Generate an executive summary DOCX.")
    parser.add_argument("--output", type=str, default="Executive_Summary_CLL.docx", help="Output DOCX filename or path")
    parser.add_argument("--date", dest="doc_date", type=str, default=date.today().isoformat(), help="Document date (YYYY-MM-DD)")
    args = parser.parse_args()

    out_path = Path(args.output).expanduser().resolve()
    out_path.parent.mkdir(parents=True, exist_ok=True)

    doc = build_document(doc_date=args.doc_date)
    doc.save(out_path)

    print(f"Wrote {out_path}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())

