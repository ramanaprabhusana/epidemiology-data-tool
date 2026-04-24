"""
Generate Word document: Epidemiology Evidence Pack – suggestions for capstone deliverable.
Run from project root: python scripts/create_suggestions_doc.py
"""
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    return p

def add_para(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    if bold:
        run.bold = True
    return p

def main():
    doc = Document()
    doc.add_heading("Epidemiology Evidence Pack - Suggestions for Capstone Deliverable", 0)
    doc.add_paragraph("PharmaACE IP Project / InsightACE Epidemiology (Beta). Recommendations to satisfy both the student IP team and the client.")
    doc.add_paragraph()

    add_heading(doc, "1. The Core Issue", level=1)
    add_para(doc, "Client side: They want trustworthy epidemiology, confidence, no AI numbers, and something that fits InsightACE. They are not giving you internal data.")
    add_para(doc, "Student side: You need a clear goal, something that feels analytical (not just a list of links), and a concrete deliverable that demonstrates expertise.")
    add_para(doc, "So the product should: use only public/source-backed data, add analytics and judgment on top of extraction, and produce client-ready artifacts (scorecards, reports, methodology).")
    doc.add_paragraph()

    add_heading(doc, "2. One \"Great Product\" Idea: Epidemiology Evidence Pack", level=1)
    add_para(doc, "Package the pipeline output into a single, repeatable product: \"Epidemiology Evidence Pack\" per indication. It is more than \"sources\": it is sourced data + analytics + documentation that PharmaACE can drop into InsightACE or use for decisions.")
    add_para(doc, "Layer 1 - Extraction (what you do now): Evidence table, reference links, tool-ready/InsightACE CSV.")
    add_para(doc, "Layer 2 - Analytics (what makes it \"great\"): Confidence scoring, source agreement, gaps/white-space, optional forecast.")
    add_para(doc, "Layer 3 - Deliverables (what impresses): One-pager per indication, KPI scorecard, methodology doc, optional dashboard.")
    add_para(doc, "That way you are not \"just\" generating sources; you are curating, validating, and explaining them.")
    doc.add_paragraph()

    add_heading(doc, "3. Concrete Ideas That Satisfy Both Sides", level=1)

    add_heading(doc, "3.1 Calculated Confidence + Rubric (high impact, analytical)", level=2)
    add_para(doc, "Idea: Replace or supplement manual high/medium/low with a rule-based confidence score (e.g. 0–100 or Low/Medium/High) from: source tier, extraction success, completeness (definition, year, geography), and optionally recency.")
    add_para(doc, "Deliverable: Confidence on each evidence row (and in exports). One-page \"Confidence scoring rubric\" (PDF/Word): what each factor means, how they combine, how to interpret.")
    add_para(doc, "Why it satisfies: Client - \"High-confidence, source-backed\" and \"confidence scored\" from the project plan. You - Clear methodology, defensible, no client data needed.")
    doc.add_paragraph()

    add_heading(doc, "3.2 KPI Scorecard That Is Truly a Scorecard (high impact)", level=2)
    add_para(doc, "Idea: Turn the KPI table into a validation scorecard: for each required metric (incidence, prevalence, etc.), show: Coverage (number of sources; best value and source; min–max range); Agreement (Do sources agree within X%?); Validation readiness (e.g. Ready / Needs review / No source based on rules).")
    add_para(doc, "Deliverable: KPI scorecard CSV/Excel (one row per metric: metric, best_value, source, range_min, range_max, n_sources, agreement_flag, validation_status). Short \"How to read the scorecard\" in the methodology doc.")
    add_para(doc, "Why it satisfies: Client - Directly addresses \"KPI scorecard (validated vs AI-generated)\". You - Clearly analytical, no internal data required.")
    doc.add_paragraph()

    add_heading(doc, "3.3 White-Space Map (high impact, strategic)", level=2)
    add_para(doc, "Idea: Automatically identify gaps: metrics × geography × year where you have no or weak evidence. Output a coverage matrix (e.g. rows = metrics, columns = sources; cell = Has value / Stub only / Missing).")
    add_para(doc, "Deliverable: White-space summary (table + one page): e.g. \"For CLL US: Incidence and prevalence covered; Stage I–III splits have limited sources; Mortality has 2 sources, need reconciliation.\" Optional: heatmap in Excel or dashboard.")
    add_para(doc, "Why it satisfies: Client - \"Clear identification of white-space gaps\" (project plan and slides). You - Shows strategic thinking without needing their data.")
    doc.add_paragraph()

    add_heading(doc, "3.4 Source Agreement & Reconciliation View (analytical)", level=2)
    add_para(doc, "Idea: For each metric that has multiple source values, produce a reconciliation view: Source A = X, Source B = Y; range; suggested \"best\" (e.g. gold source or median) with a short rule.")
    add_para(doc, "Deliverable: Reconciliation table (metric, geography, year, value_by_source, range, recommended_value, note). Optional: \"Reconciliation rules\" one-pager.")
    add_para(doc, "Why it satisfies: Client - \"Validate, reconcile, and standardize\" and \"assumption-clear\". You - Clear analytical step beyond extraction.")
    doc.add_paragraph()

    add_heading(doc, "3.5 One-Pager \"Evidence Summary\" Per Indication (impressive, low effort)", level=2)
    add_para(doc, "Idea: For each of the 5 indications, one single-page summary: key numbers (incidence, prevalence, best source, year); confidence; top 3–5 sources; main gaps (white-space); one-line caveats.")
    add_para(doc, "Deliverable: One PDF per indication (e.g. CLL_US_Evidence_Summary.pdf).")
    add_para(doc, "Why it satisfies: Client - Something they can share internally. You - Demonstrates communication and synthesis.")
    doc.add_paragraph()

    add_heading(doc, "3.6 Lightweight Forecast with Guardrails (optional, very analytical)", level=2)
    add_para(doc, "Idea: Use only your extracted incidence/prevalence to run a simple trend forecast (e.g. linear or exponential growth) with low/base/high scenarios. Always ship with an \"Assumptions and limitations\" note.")
    add_para(doc, "Deliverable: Forecast table (year, low/base/high) per metric; one-page \"Forecast assumptions and limitations.\"")
    add_para(doc, "Why it satisfies: Client - Connects to \"Forecasting\" in InsightACE. You - Clearly analytical and reusable.")
    doc.add_paragraph()

    add_heading(doc, "3.7 Methodology + Reproducibility Pack (recognition and handoff)", level=2)
    add_para(doc, "Idea: One short methodology document (5–7 pages): how sources are selected; how extraction works; how confidence is calculated; how the KPI scorecard and white-space are produced. Plus a reproducibility pack: config files, how to run for a new indication.")
    add_para(doc, "Deliverable: \"Epidemiology Evidence Pack – Methodology and reproducibility\" (PDF + repo/config snapshot).")
    add_para(doc, "Why it satisfies: Client - \"Reusable, transparent pipelines\" and easier handoff. You - Positions the team as rigorous.")
    doc.add_paragraph()

    add_heading(doc, "4. How to Package It for Both Sides", level=1)
    add_para(doc, "Name the product: e.g. \"Epidemiology Evidence Pack\" or \"Source-backed epidemiology pack for InsightACE\".", bold=True)
    add_para(doc, "One-sentence for PharmaACE: \"We deliver, per indication, a curated evidence table, a confidence-scored and source-mapped dataset, a KPI scorecard (validated vs. AI-ready), a white-space summary, optional reconciliation view and forecast, and a methodology + reproducibility pack-all from public sources, no AI-generated numbers, ready to feed InsightACE or your own analytics.\"")
    add_para(doc, "One-sentence for your team: \"We do not just list sources; we extract, score, reconcile, gap-analyze, and document so the deliverable is analytical and client-ready.\"")
    doc.add_paragraph()

    add_heading(doc, "5. Suggested Priority Order", level=1)
    add_para(doc, "Must-have (core \"great product\"): Calculated confidence + rubric; KPI scorecard (coverage, agreement, validation status); White-space summary (gaps per metric/indication); One-pager per indication (evidence summary); Methodology + reproducibility doc.")
    add_para(doc, "Strong add-on: Source agreement / reconciliation table (and short rules).")
    add_para(doc, "If time allows: Lightweight forecast (low/base/high + assumptions note); Dashboard (coverage heatmap, scorecard view) using your existing export.")
    doc.add_paragraph()

    add_heading(doc, "6. How This Addresses \"No Clear Goal\" and \"Not Analytical\"", level=1)
    add_para(doc, "Goal becomes: \"Deliver the Epidemiology Evidence Pack per indication: evidence + confidence + scorecard + white-space + methodology (and optionally reconciliation + forecast).\"")
    add_para(doc, "Analytical part is explicit: scoring, agreement, gaps, reconciliation, optional forecasting, and a clear methodology-all on top of extraction.")
    add_para(doc, "Recognition: You are not \"just\" building a scraper; you are defining how epidemiology should be sourced, scored, and validated for InsightACE, which matches the project plan and the slides.")
    doc.add_paragraph()

    doc.add_paragraph("- End of document -")
    out_path = Path(__file__).resolve().parents[1] / "Epidemiology_Evidence_Pack_Suggestions.docx"
    doc.save(str(out_path))
    print(f"Saved: {out_path}")
    return out_path

if __name__ == "__main__":
    main()
