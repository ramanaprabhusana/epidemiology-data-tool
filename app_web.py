"""
Web page alternative to the Streamlit app.
Run with: python app_web.py  then open http://localhost:5000
Same flow: choose indication + country → Get data → use output files.
"""

import sys
from pathlib import Path

ROOT = Path(__file__).resolve().parent
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))

import io
import zipfile
import yaml
from flask import Flask, render_template, request, jsonify, send_file

from src.pipeline.runner import run_pipeline


def _json_safe(obj):
    """Convert to JSON-serializable types (e.g. numpy int -> int)."""
    if obj is None or isinstance(obj, (bool, str, int, float)):
        return obj
    if hasattr(obj, "item"):
        return obj.item()
    if isinstance(obj, (list, tuple)):
        return [_json_safe(x) for x in obj]
    if isinstance(obj, dict):
        return {str(k): _json_safe(v) for k, v in obj.items()}
    return str(obj)

app = Flask(
    __name__,
    template_folder=str(ROOT / "templates_web"),
    static_folder=str(ROOT / "static_web"),
)
app.config["MAX_CONTENT_LENGTH"] = 16 * 1024 * 1024  # 16 MB for uploads
OUTPUT_DIR = ROOT / "output"


def load_ui_options():
    config_path = ROOT / "config" / "ui_options.yaml"
    default = {"indications": [{"id": "cll", "label": "CLL"}], "countries": [{"id": "US", "label": "United States"}]}
    if not config_path.exists():
        return default
    try:
        with open(config_path, "r") as f:
            data = yaml.safe_load(f)
        return data if isinstance(data, dict) else default
    except Exception:
        return default


@app.route("/")
def index():
    options = load_ui_options()
    return render_template("index.html", indications=options.get("indications", []), countries=options.get("countries", []))


@app.route("/run", methods=["POST"])
def run():
    try:
        data = request.get_json(silent=True) or {}
    except Exception:
        data = {}
    indication = (data.get("indication") or "").strip() or "CLL"
    country = (data.get("country") or "").strip() or "US"
    export_dashboard = data.get("export_dashboard", True)
    # PubMed is always on; client cannot disable it.
    use_pubmed = True
    add_pubmed_stubs = True
    evidence_path = None
    if data.get("evidence_filename"):
        up = ROOT / "output" / "temp_upload" / (data.get("evidence_filename") or "").lstrip("/")
        if up.exists() and up.suffix.lower() == ".csv":
            evidence_path = up
    try:
        result = run_pipeline(
            indication=indication,
            country=country,
            evidence_path=evidence_path,
            config_dir=ROOT / "config",
            output_dir=OUTPUT_DIR,
            export_dashboard=export_dashboard,
            include_forecast=True,
            use_pubmed=use_pubmed,
            add_pubmed_stubs=add_pubmed_stubs,
            max_run_seconds=45 * 60,  # 45 min cap; priority sources get deep-dive first
        )
    except Exception as e:
        return jsonify({
            "success": False,
            "message": str(e),
            "output_files": [],
            "record_count": 0,
            "sources_explored": [],
            "kpi_labels": [],
        })

    try:
        # Build JSON-serializable response (no DataFrames, no numpy types)
        paths = result.get("paths", {}) or {}
        file_list = []
        for key, full_path in paths.items():
            if key == "dashboard_dir":
                continue
            try:
                p = Path(str(full_path))
                if p.is_file() and OUTPUT_DIR in p.parents:
                    file_list.append(p.name)
            except Exception:
                pass
        # Deduplicate by name; put Excel (.xlsx) first
        seen = set()
        dedup = [n for n in file_list if n not in seen and not seen.add(n)]
        output_files_dedup = sorted(dedup, key=lambda n: (0 if n.lower().endswith(".xlsx") else 1, n))
        response = {
            "success": bool(result.get("success", False)),
            "message": str(result.get("message", "")),
            "record_count": int(result.get("record_count", 0)),
            "output_files": output_files_dedup,
            "validation_passed": bool(result.get("validation_passed", True)),
            "validation_report": _json_safe(result.get("validation_report", [])),
            "sources_explored": list(result.get("sources_explored", [])),
            "kpi_labels": list(result.get("kpi_labels", [])),
        }
        return jsonify(response)
    except Exception as e:
        return jsonify({
            "success": False,
            "message": f"Response error: {e}",
            "record_count": 0,
            "output_files": [],
            "sources_explored": [],
            "kpi_labels": [],
        })


ALLOWED_DOWNLOAD_EXT = (".csv", ".xlsx", ".db", ".md")


@app.route("/download/<filename>")
def download(filename):
    """Serve a file from output/ if it exists and is an allowed type."""
    if not filename or ".." in filename or "/" in filename or "\\" in filename:
        return "Invalid file", 400
    if not any(filename.lower().endswith(ext) for ext in ALLOWED_DOWNLOAD_EXT):
        return "Invalid file type", 400
    path = OUTPUT_DIR / filename
    if not path.exists() or not path.is_file():
        return "File not found", 404
    try:
        return send_file(path, as_attachment=True, download_name=filename)
    except Exception:
        return "Error sending file", 500


@app.route("/download_zip", methods=["POST"])
def download_zip():
    """
    Build a zip of selected (or all) output files and return it.
    Expects JSON body: { "filenames": ["a.csv", "b.csv"] }. If filenames is empty or missing, zips all files in output/.
    """
    try:
        data = request.get_json(silent=True) or {}
        filenames = data.get("filenames")
        if filenames is not None and not isinstance(filenames, list):
            filenames = []
    except Exception:
        filenames = []

    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
        if filenames:
            for name in filenames:
                name = name.strip() if isinstance(name, str) else str(name).strip()
                if not name or ".." in name or "/" in name or "\\" in name:
                    continue
                if not any(name.lower().endswith(ext) for ext in ALLOWED_DOWNLOAD_EXT):
                    continue
                path = OUTPUT_DIR / name
                if path.is_file() and path.exists() and OUTPUT_DIR.resolve() in path.resolve().parents:
                    zf.write(path, name)
        else:
            # Zip all allowed files in output dir (non-recursive: top-level only)
            for path in sorted(OUTPUT_DIR.iterdir()):
                if path.is_file() and any(path.name.lower().endswith(ext) for ext in ALLOWED_DOWNLOAD_EXT):
                    zf.write(path, path.name)
        # Include dashboard/ contents if present
        dashboard_dir = OUTPUT_DIR / "dashboard"
        if dashboard_dir.is_dir():
            for path in sorted(dashboard_dir.iterdir()):
                if path.is_file() and any(path.name.lower().endswith(ext) for ext in ALLOWED_DOWNLOAD_EXT):
                    zf.write(path, f"dashboard/{path.name}")

    buf.seek(0)
    download_name = "epidemiology_outputs.zip"
    return send_file(buf, as_attachment=True, download_name=download_name, mimetype="application/zip")


ALLOWED_UPLOAD_EXT = (".csv", ".xlsx", ".xls", ".docx")
EVIDENCE_CSV_COLUMNS = (
    "indication", "metric", "value", "source_citation", "source_tier", "definition",
    "population", "year_or_range", "geography", "split_logic", "source_url", "notes", "confidence",
)


def _excel_to_evidence_csv(path: Path, out_csv: Path) -> None:
    """Read first sheet of Excel file and save as evidence CSV."""
    import pandas as pd
    df = pd.read_excel(path, sheet_name=0)
    # Normalize column names to match evidence schema (strip whitespace, case-insensitive match)
    col_map = {c.strip().lower(): c for c in df.columns if isinstance(c, str)}
    rename = {}
    for want in EVIDENCE_CSV_COLUMNS:
        for k, orig in col_map.items():
            if k == want.lower().replace(" ", "_") or k == want.lower():
                rename[orig] = want
                break
    if rename:
        df = df.rename(columns=rename)
    df.to_csv(out_csv, index=False, encoding="utf-8")


def _docx_to_evidence_csv(path: Path, out_csv: Path) -> None:
    """Extract first table from Word document and save as evidence CSV."""
    try:
        from docx import Document
    except ImportError:
        raise ValueError("python-docx is required for Word uploads; pip install python-docx")
    import pandas as pd
    doc = Document(path)
    if not doc.tables:
        # No table: create minimal CSV with headers so pipeline doesn't fail
        pd.DataFrame(columns=EVIDENCE_CSV_COLUMNS).to_csv(out_csv, index=False, encoding="utf-8")
        return
    table = doc.tables[0]
    rows = []
    ncols = 0
    for row in table.rows:
        cells = [cell.text.strip() if cell.text else "" for cell in row.cells]
        ncols = max(ncols, len(cells))
        rows.append(cells)
    for i in range(len(rows)):
        rows[i] = (rows[i] + [""] * ncols)[:ncols]
    if not rows:
        pd.DataFrame(columns=EVIDENCE_CSV_COLUMNS).to_csv(out_csv, index=False, encoding="utf-8")
        return
    headers = rows[0]
    df = pd.DataFrame(rows[1:], columns=headers if any(h for h in headers) else range(len(headers)))
    if df.empty:
        df = pd.DataFrame(columns=EVIDENCE_CSV_COLUMNS)
    else:
        col_map = {str(c).strip().lower().replace(" ", "_").replace("-", "_"): c for c in df.columns}
        rename = {}
        for want in EVIDENCE_CSV_COLUMNS:
            w = want.lower().replace(" ", "_").replace("-", "_")
            if w in col_map:
                rename[col_map[w]] = want
        if rename:
            df = df.rename(columns=rename)
    df.to_csv(out_csv, index=False, encoding="utf-8")


@app.route("/upload", methods=["POST"])
def upload():
    """Accept evidence upload (CSV, Excel, or Word); save as CSV in temp_upload and return filename."""
    if "file" not in request.files:
        return jsonify({"ok": False, "error": "No file"}), 400
    f = request.files["file"]
    if not f.filename:
        return jsonify({"ok": False, "error": "No file selected"}), 400
    ext = Path(f.filename).suffix.lower()
    if ext not in ALLOWED_UPLOAD_EXT:
        return jsonify({"ok": False, "error": "Use CSV, Excel (.xlsx, .xls), or Word (.docx)"}), 400
    temp_dir = OUTPUT_DIR / "temp_upload"
    temp_dir.mkdir(parents=True, exist_ok=True)
    base = "upload_" + Path(f.filename).stem.replace(" ", "_")[:60]
    safe_name = base + ".csv"
    path = temp_dir / (base + ext)
    out_csv = temp_dir / safe_name
    try:
        if ext == ".csv":
            f.save(str(out_csv))
        else:
            f.save(str(path))
        if ext in (".xlsx", ".xls"):
            _excel_to_evidence_csv(path, out_csv)
            path.unlink(missing_ok=True)
        elif ext == ".docx":
            _docx_to_evidence_csv(path, out_csv)
            path.unlink(missing_ok=True)
    except Exception as e:
        if path.exists():
            path.unlink(missing_ok=True)
        return jsonify({"ok": False, "error": str(e)}), 400
    return jsonify({"ok": True, "filename": safe_name})


if __name__ == "__main__":
    import socket
    import threading
    import time
    import webbrowser
    # use_reloader=False so we don't get extra browser tabs when the reloader restarts the process.
    port = 5000
    for _ in range(10):
        try:
            sock = socket.socket(socket.AF_INET, socket.SOCK_STREAM)
            sock.bind(("127.0.0.1", port))
            sock.close()
            break
        except OSError:
            port += 1
    else:
        port = 5000  # fallback; will show error when app.run runs
    port_file = ROOT / ".webapp_port"
    try:
        port_file.write_text(str(port))
    except Exception:
        pass
    url = "http://127.0.0.1:{}".format(port)
    print("Epidemiology Data Tool — Web")
    print("Open in browser: {}".format(url))

    def _open_browser_once():
        for _ in range(30):
            time.sleep(1)
            try:
                import urllib.request
                urllib.request.urlopen(url, timeout=2)
                break
            except Exception:
                continue
        try:
            webbrowser.open(url)
        except Exception:
            pass

    threading.Thread(target=_open_browser_once, daemon=True).start()
    app.run(host="127.0.0.1", port=port, debug=True, use_reloader=False)
